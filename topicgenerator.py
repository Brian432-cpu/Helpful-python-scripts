import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import quote
import nltk
from nltk.tokenize import sent_tokenize
import re
from concurrent.futures import ThreadPoolExecutor, as_completed

# Make sure you have nltk data
nltk.download('punkt', quiet=True)

def duckduckgo_search(query, num_results=5):
    print(f"[*] Searching DuckDuckGo for: {query}")
    url = f"https://duckduckgo.com/html/?q={quote(query)}"
    headers = {'User-Agent': 'Mozilla/5.0'}
    res = requests.get(url, headers=headers, timeout=10)
    soup = BeautifulSoup(res.text, 'html.parser')
    links = []
    for a in soup.select('a.result__a', limit=num_results):
        href = a.get('href')
        if href and href.startswith('http'):
            links.append(href)
    return links

def extract_text_from_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        res = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(res.text, 'html.parser')
        paragraphs = soup.find_all('p')
        text = ' '.join(p.get_text() for p in paragraphs)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    except Exception as e:
        print(f"[!] Failed to extract text from {url}: {e}")
        return ''

def summarize_text(text, num_sentences=10):
    sentences = sent_tokenize(text)
    if len(sentences) <= num_sentences:
        return ' '.join(sentences)

    # Simple frequency-based summarization
    words = re.findall(r'\w+', text.lower())
    freq = {}
    for word in words:
        freq[word] = freq.get(word, 0) + 1

    sentence_scores = {}
    for sentence in sentences:
        for word in nltk.word_tokenize(sentence.lower()):
            if word in freq:
                sentence_scores[sentence] = sentence_scores.get(sentence, 0) + freq[word]

    ranked = sorted(sentence_scores, key=sentence_scores.get, reverse=True)
    summary_sentences = ranked[:num_sentences]
    summary = ' '.join(summary_sentences)
    return summary

def save_to_word(topic, summary, urls):
    doc = Document()
    doc.add_heading(f'Summary Report: {topic}', level=1)

    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)

    doc.add_heading('Sources', level=2)
    for url in urls:
        doc.add_paragraph(url, style='List Bullet')

    filename = f'summary_{topic.replace(" ", "_")}.docx'
    doc.save(filename)
    print(f'[+] Summary saved to {filename}')

def fetch_all_texts(urls):
    all_texts = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        future_to_url = {executor.submit(extract_text_from_url, url): url for url in urls}
        for future in as_completed(future_to_url):
            url = future_to_url[future]
            try:
                text = future.result()
                if text:
                    all_texts.append(text)
            except Exception as e:
                print(f"[!] Error fetching {url}: {e}")
    return all_texts

def main():
    topic = input('Enter a topic to summarize: ').strip()
    urls = duckduckgo_search(topic)

    if not urls:
        print('[!] No URLs found for the topic.')
        return

    print('[*] Fetching contents from all URLs in parallel...')
    all_texts = fetch_all_texts(urls)

    combined_text = ' '.join(all_texts)
    if not combined_text:
        print('[!] No text extracted from sources.')
        return

    print('[*] Generating summary...')
    summary = summarize_text(combined_text, num_sentences=10)

    save_to_word(topic, summary, urls)

if __name__ == '__main__':
    main()